from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import textwrap

input_md = 'documentation.md'
output_docx = 'documentation.docx'

with open(input_md, 'r', encoding='utf-8') as f:
    lines = f.read().splitlines()

doc = Document()

# Set base font and sizes for Normal and Headings to ensure consistent language/typography
normal_style = doc.styles['Normal']
normal_font = normal_style.font
normal_font.name = 'Times New Roman'
normal_font.size = Pt(12)

try:
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
except Exception:
    # If headings aren't present in the style collection, continue gracefully
    pass

# Simple Markdown-like parser: headings and paragraphs
for line in lines:
    if line.startswith('# '):
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith('---'):
        # horizontal rule -> blank paragraph for separation
        doc.add_paragraph('')
    elif line.strip() == '':
        doc.add_paragraph('')
    else:
        wrapped = '\n'.join(textwrap.wrap(line, width=110))
        p = doc.add_paragraph(wrapped)
        p.style = doc.styles['Normal']

# Add centered page numbers (Arabic numerals) to the footer of every section
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH

for section in doc.sections:
    footer = section.footer
    # Clear existing footer paragraphs
    for p in list(footer.paragraphs):
        # remove text but keep paragraph objects
        p.clear()
    # Add centered page number field
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run()
    # Insert a PAGE field; formatting (Roman vs Arabic) is handled by Word; use Arabic here
    try:
        run._r.add_fldSimple('PAGE')
    except Exception:
        # If fldSimple isn't available, fallback to a literal page placeholder
        run.text = 'Page'

# Save
doc.save(output_docx)
print('DOCX generated:', output_docx)
